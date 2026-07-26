import os
import re
import pdfplumber
import pypdf
import docx

# Optional spaCy loader with fallback to regex
spacy_nlp = None
try:
    import spacy
    try:
        spacy_nlp = spacy.load("en_core_web_sm")
    except Exception:
        spacy_nlp = None
except Exception:
    spacy_nlp = None

# Comprehensive dictionary of technical & soft skills
TECHNICAL_SKILLS = [
    # Programming Languages
    "Python", "Java", "C++", "C#", "C", "JavaScript", "TypeScript", "PHP", "Ruby", "Go", "Rust", "Kotlin", "Swift", "R", "Scala",
    
    # Web & Frameworks
    "HTML", "CSS", "Bootstrap", "Tailwind", "Flask", "Django", "FastAPI", "React", "Angular", "Vue", "Node.js", "Express", "REST API", "GraphQL",
    
    # Databases & Cloud
    "SQL", "MySQL", "PostgreSQL", "SQLite", "MongoDB", "Redis", "Firebase", "AWS", "Azure", "GCP", "Docker", "Kubernetes", "Git", "GitHub", "Linux",
    
    # Data Science & AI/ML
    "Machine Learning", "Deep Learning", "Data Science", "AI", "Artificial Intelligence", "Scikit-Learn", "TensorFlow", "PyTorch", 
    "Pandas", "NumPy", "Matplotlib", "Seaborn", "NLP", "Natural Language Processing", "Computer Vision", "OpenCV", "Tableau", "Power BI"
]

def extract_text_from_pdf(file_path):
    text = ""
    # Method 1: pdfplumber
    try:
        with pdfplumber.open(file_path) as pdf:
            for page in pdf.pages:
                page_text = page.extract_text()
                if page_text:
                    text += page_text + "\n"
    except Exception as e:
        print(f"[PDF Extraction Warning] pdfplumber failed: {e}. Trying pypdf...")
        text = ""

    # Fallback Method 2: pypdf
    if not text.strip():
        try:
            reader = pypdf.PdfReader(file_path)
            for page in reader.pages:
                page_text = page.extract_text()
                if page_text:
                    text += page_text + "\n"
        except Exception as e:
            print(f"[PDF Extraction Error] pypdf failed: {e}")
            
    return text.strip()

def extract_text_from_docx(file_path):
    text = ""
    try:
        doc = docx.Document(file_path)
        for paragraph in doc.paragraphs:
            if paragraph.text:
                text += paragraph.text + "\n"
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text += cell.text + " "
                text += "\n"
    except Exception as e:
        print(f"[DOCX Extraction Error] {e}")
    return text.strip()

def extract_resume_text(file_path):
    ext = os.path.splitext(file_path)[1].lower()
    if ext == '.pdf':
        return extract_text_from_pdf(file_path)
    elif ext in ['.docx', '.doc']:
        return extract_text_from_docx(file_path)
    else:
        return ""

def extract_skills(text):
    if not text:
        return []
    
    extracted = set()
    cleaned_text = " " + text.lower() + " "
    
    # 1. Regex Matcher for defined Technical Skills
    for skill in TECHNICAL_SKILLS:
        skill_lower = skill.lower()
        escaped = re.escape(skill_lower)
        pattern = r'(?:^|[\s,./\(\)\[\]:]+)' + escaped + r'(?:$|[\s,./\(\)\[\]:]+)'
        if re.search(pattern, cleaned_text):
            extracted.add(skill)

    # 2. spaCy Named Entity / Noun Chunk enhancement if spaCy is loaded
    if spacy_nlp:
        try:
            doc = spacy_nlp(text[:5000])  # limit length
            for chunk in doc.noun_chunks:
                chunk_text = chunk.text.strip().title()
                for skill in TECHNICAL_SKILLS:
                    if skill.lower() == chunk_text.lower():
                        extracted.add(skill)
        except Exception:
            pass

    return sorted(list(extracted))
